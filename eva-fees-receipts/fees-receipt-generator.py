import mysql.connector
from docx import Document
from num2words import num2words
import os
from datetime import datetime
from decimal import Decimal
from docx2pdf import convert
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

def send_email_with_attachment(recipient_email, student_name, pdf_filepath, invoice_number, smtp_config):
    msg = MIMEMultipart()
    msg['From'] = f"{smtp_config['sender_name']} <{smtp_config['sender_email']}>"
    msg['To'] = recipient_email
    msg['Subject'] = f"Fee Receipt from VisionAstraa EV Academy - {invoice_number}"

    body = f"""
Dear {student_name},

Please find your fee receipt attached to this email.

If you have any questions, feel free to contact us.

Best regards,
VisionAstraa EV Academy
    """
    msg.attach(MIMEText(body, 'plain'))

    try:
        with open(pdf_filepath, "rb") as attachment:
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(attachment.read())
        encoders.encode_base64(part)
        part.add_header(
            'Content-Disposition',
            f"attachment; filename= {os.path.basename(pdf_filepath)}",
        )
        msg.attach(part)
    except IOError as e:
        print(f"  - ERROR: Could not read attachment file for email. Error: {e}")
        return 
    
    server = None
    try:
        server = smtplib.SMTP_SSL(smtp_config['host'], smtp_config['port'])
        server.login(smtp_config['sender_email'], smtp_config['password'])
        server.sendmail(smtp_config['sender_email'], recipient_email, msg.as_string())
    finally:
        if server:
            server.quit()


def generate_receipts():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    
    smtp_config = {
        'host': 'smtp.hostinger.com',
        'port': 465,
        'sender_email': 'careers@visionastraa.in',
        'sender_name': 'VisionAstraa EV Academy',
        'password': 'Z1SIOO0A9b~'
    }

    online_template_path = os.path.join(script_dir, "EVA_Fee_Receipt_Template_Online.docx")
    cash_template_path = os.path.join(script_dir, "EVA_Fee_Receipt_Template_Cash.docx")

    if not os.path.exists(online_template_path):
        print(f"Error: '{os.path.basename(online_template_path)}' not found in {script_dir}")
        return
    if not os.path.exists(cash_template_path):
        print(f"Error: '{os.path.basename(cash_template_path)}' not found in {script_dir}")
        return

    connection = None
    try:
        connection = mysql.connector.connect(
            host="srv1640.hstgr.io",
            user="u707137586_EV_Fees",
            password="@1Qk?gYQ>Ioj",
            database="u707137586_EV_Fees"
        )

        if connection.is_connected():
            print("Successfully connected to the database.")

        cursor = connection.cursor(dictionary=True)

        cursor.execute("SELECT id, full_name, email, phone_number, program_category, base_rate, payment_mode, invoice_date FROM students WHERE payment_status = 'pending'")
        students = cursor.fetchall()

        if not students:
            print("No students found with 'pending' payment status. Exiting.")
            return

        output_dir = os.path.join(script_dir, "generated_receipts")
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
            print(f"Created directory: {output_dir}")

        for student in students:
            if not student['invoice_date']:
                print(f"\n- WARNING: Skipping {student['full_name']} (ID: {student['id']}) because 'invoice_date' is not set in the database.")
                continue

            payment_mode = student.get('payment_mode', 'Online').strip().title()
            print(f"\nProcessing {student['full_name']} (ID: {student['id']}, Payment Mode: {payment_mode})...")

            invoice_number = f"FEE-25B01-{str(student['id']).zfill(3)}"

            invoice_date = student['invoice_date'].strftime("%d/%m/%Y")

            if payment_mode == 'Online':
                doc = Document(online_template_path)
                rate = student['base_rate']
                gst = rate * Decimal('0.18')
                total = rate + gst
                total_in_words = f"{num2words(total, lang='en_IN').title()} Paid"
                
                replacements = {
                    "[RATE]": f"{rate:,.2f}",
                    "[GST]": f"{gst:,.2f}",
                    "[TOTAL]": f"{total:,.2f}",
                }
            elif payment_mode == 'Cash':
                doc = Document(cash_template_path)
                rate = student['base_rate']
                total = rate
                total_in_words = f"{num2words(total, lang='en_IN').title()} Paid"

                replacements = {
                    "[RATE]": f"{rate:,.2f}",
                    "[TOTAL]": f"{total:,.2f}",
                }
            else:
                print(f"  - WARNING: Unknown payment mode '{payment_mode}' for {student['full_name']}. Skipping.")
                continue
            
            common_replacements = {
                "FEE-25B01-00X": invoice_number,
                "0X/0X/2025": invoice_date,
                "[STUDENT FULL NAME]": student['full_name'],
                "[EMAIL]": student['email'],
                "[PHONE]": student['phone_number'],
                "[CATEGORY]": student['program_category'],
                "One Lakh and Three Hundred [MENTION “PAID”]": total_in_words,
            }
            replacements.update(common_replacements)

            for para in doc.paragraphs:
                for key, value in replacements.items():
                    if key in para.text:
                        para.text = para.text.replace(key, str(value))

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in replacements.items():
                            if key in cell.text:
                                cell.text = cell.text.replace(key, str(value))
            
            print("  - Replaced placeholders.")

            docx_filename = os.path.join(output_dir, f"Fee_Receipt_{student['full_name']}_{invoice_number}.docx")
            doc.save(docx_filename)
            print(f"  -> Successfully generated DOCX: {os.path.basename(docx_filename)}")

            pdf_filename = os.path.join(output_dir, f"Fee_Receipt_{student['full_name']}_{invoice_number}.pdf")
            try:
                print(f"  - Converting to PDF...")
                convert(docx_filename, pdf_filename)
                print(f"  -> Successfully converted to PDF: {os.path.basename(pdf_filename)}")
            except Exception as e:
                print(f"  - ERROR: Failed to convert to PDF. Please ensure Microsoft Word or LibreOffice is installed. Error: {e}")
                continue
            
            try:
                print("  - Sending email with receipt...")
                send_email_with_attachment(
                    student['email'],
                    student['full_name'],
                    pdf_filename,
                    invoice_number,
                    smtp_config
                )
                print(f"  -> Email sent successfully to {student['email']}.")
            except Exception as e:
                print(f"  - ERROR: Failed to send email. Error: {e}")
            
            pdf_data = None
            try:
                with open(pdf_filename, 'rb') as f:
                    pdf_data = f.read()
                print("  - Read PDF file content for database storage.")
            except IOError as e:
                print(f"  - ERROR: Could not read PDF file. Error: {e}")
                continue

            if pdf_data:
                try:
                    update_query = "UPDATE students SET payment_status = 'paid', invoice_number = %s, receipt_pdf = %s WHERE id = %s"
                    cursor.execute(update_query, (invoice_number, pdf_data, student['id']))
                    connection.commit()
                    print(f"  - Updated payment status and saved receipt for {student['full_name']} in the database.")
                except mysql.connector.Error as err:
                    print(f"  - ERROR: Failed to update database with PDF. Error: {err}")
            
            try:
                os.remove(docx_filename)
                os.remove(pdf_filename)
                print("  - Cleaned up local files.")
            except OSError as e:
                print(f"  - WARNING: Could not clean up local files. Error: {e}")

    except mysql.connector.Error as err:
        print(f"Error connecting to database: {err}")
    finally:
        if connection and connection.is_connected():
            cursor.close()
            connection.close()
            print("\nDatabase connection closed.")

if __name__ == "__main__":
    generate_receipts()
